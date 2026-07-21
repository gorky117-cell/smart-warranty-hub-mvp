from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Smart_Warranty_Hub_Hackathon_Demo_Guide.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)


def set_run_font(run, size, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_style(style, size, color, before, after, line_spacing=1.25, bold=False):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def add_heading(doc, text, level=1):
    style = {1: "Heading 1", 2: "Heading 2"}[level]
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_list(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], 11, INK, 0, 6)
    set_style(doc.styles["Heading 1"], 16, BLUE, 18, 10, bold=True)
    set_style(doc.styles["Heading 2"], 13, BLUE, 14, 7, bold=True)
    set_style(doc.styles["List Bullet"], 11, INK, 0, 4)
    set_style(doc.styles["List Number"], 11, INK, 0, 4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("Smart Warranty Hub | Hackathon Demo Guide"), 9, DARK_BLUE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Controlled demo and judge presentation reference"), 9, DARK_BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("Smart Warranty Hub"), 24, DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("Hackathon Demo Guide"), 14, DARK_BLUE)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(12)
    set_run_font(lead.add_run("Pitch: "), 11, INK, bold=True)
    set_run_font(
        lead.add_run(
            "Smart Warranty Hub turns a customer's invoice into clear warranty guidance, proactive care and risk support, while giving OEM and TPA teams privacy-safe signals to act earlier."
        ),
        11,
    )

    add_heading(doc, "Two-Minute Judge Flow")
    add_list(
        doc,
        [
            "Start on the public site and introduce Smart Warranty Hub as a post-purchase platform for customers, OEMs and TPAs.",
            "Sign in and open the Neo dashboard.",
            "Upload a sample invoice or choose an available warranty.",
            "Show the warranty summary: product details, coverage, exclusions, claim steps, confidence and evidence/source context.",
            "Show predictive risk, behaviour questions and preventative care or expiry guidance.",
            "Open the Resolution checklist. Explain that it is draft-only and cannot submit a claim, contact an OEM or run a device action.",
            "Switch to the OEM dashboard. Show aggregate insight, source verification and controlled question/recommendation workflows.",
            "Close with the privacy boundary: direct OEM sharing needs explicit consent, and aggregate telemetry is cohort-suppressed.",
        ],
        numbered=True,
    )

    add_heading(doc, "What Makes the Demo Credible")
    add_list(
        doc,
        [
            "Invoice upload is size/type restricted and stored with server-generated filenames.",
            "Warranty access is protected by authenticated ownership checks.",
            "The OpenAI lane is optional and safely falls back when unavailable.",
            "Telemetry strips direct identifiers before use and OEM aggregates require a minimum cohort.",
            "High-risk/cost paths have CSRF protection, rate limits, request IDs and per-user AI quotas.",
            "The warranty-resolution agent remains draft-only and records an audit trace.",
            "The full automated suite currently passes: 122 tests.",
        ],
    )

    add_heading(doc, "How to Describe Metrics")
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade_cell(cell, "E8EEF5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run("Use this wording: "), 11, DARK_BLUE, bold=True)
    set_run_font(
        paragraph.add_run(
            "The repository includes controlled 50-case evaluations for the OCR, predictive, nudge, service, OEM dispatch, watchdog and remediation workflows. These are test metrics, not live customer results."
        ),
        11,
        DARK_BLUE,
    )
    warning = doc.add_paragraph()
    warning.paragraph_format.space_before = Pt(8)
    set_run_font(warning.add_run("Do not claim live cost, turnaround, failure-prevention or OEM-outcome improvements without independently measured customer data."), 11, RGBColor(122, 90, 0), bold=True)

    add_heading(doc, "Judge Questions")
    answers = [
        ("Is this production-ready?", "It is ready for a controlled hackathon demo and pilot. A full production rollout would move local/process-level stores to managed Postgres, object storage and shared rate-limit/quota infrastructure."),
        ("What does AI decide?", "AI enriches and explains information, but deterministic warranty data, policy checks, consent, ownership and human-review boundaries remain authoritative. The agent cannot execute claims or device actions."),
        ("How is customer data protected?", "Users are isolated by ownership checks. Telemetry is sanitized, aggregate OEM insight has cohort suppression, and direct OEM communication requires separate explicit consent."),
        ("What is the business value?", "Customers get clearer coverage and earlier care guidance; OEM and TPA teams get privacy-safe early signals for support, service and product-quality workflows."),
    ]
    for question, answer in answers:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        set_run_font(paragraph.add_run(question + " "), 11, DARK_BLUE, bold=True)
        set_run_font(paragraph.add_run(answer), 11)

    add_heading(doc, "Presentation Checklist")
    add_list(
        doc,
        [
            "Use the active master branch for any code walkthrough.",
            "Set master as the GitHub default branch before judges review the repository.",
            "Keep one sample invoice and one known login ready.",
            "Demonstrate the happy path; do not rely on an external OCR/LLM call during the presentation.",
            "If an optional provider is unavailable, explain that deterministic fallbacks preserve the core flow.",
        ],
    )

    doc.core_properties.title = "Smart Warranty Hub Hackathon Demo Guide"
    doc.core_properties.subject = "Controlled demo and judge presentation guide"
    doc.core_properties.author = "Smart Warranty Hub"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
