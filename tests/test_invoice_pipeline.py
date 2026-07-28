import os
from pathlib import Path
from datetime import datetime

from fastapi.testclient import TestClient
from fpdf import FPDF

from app.main import app
from app.db import SessionLocal
from app.db_models import PipelineJobDB, WarrantyDB, WarrantySummaryDB, ParsedFieldDB
from app.services import invoice_pipeline, summary_engine, terms_lookup
from app.services.openai_intelligence import merge_invoice_enrichment
from app.services.ingestion import extract_product_fields, ingest_artifact
from app.services.canonical import canonicalize_artifact
from app.services.warranty_discovery import discover_sources
from app.services.terms_lookup import lookup_terms
from app.models import ArtifactType, CanonicalWarranty, TermsResult
from app.services.warranty_discovery import DiscoverySource
from app.services.warranty_parser import ParsedTerms


def _make_pdf(path: Path, text: str) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, text)
    pdf.output(str(path))


def test_upload_creates_job(tmp_path):
    client = TestClient(app)
    login = client.post(
        "/auth/login",
        data={"username": "admin", "password": "admin123"},
        headers={"accept": "application/json"},
    )
    assert login.status_code == 200
    token = login.json().get("access_token")
    assert token

    sample_path = tmp_path / "invoice.txt"
    sample_path.write_text("Brand: Acmeco Model: ZX-100 Purchase date: 2025-01-01", encoding="utf-8")
    with sample_path.open("rb") as fh:
        resp = client.post(
            "/artifacts/upload",
            files={"file": ("invoice.txt", fh, "text/plain")},
            data={"type": "invoice"},
            headers={"Authorization": f"Bearer {token}"},
        )
    assert resp.status_code == 200
    payload = resp.json()
    assert payload.get("job_id")
    assert payload.get("warranty_id")
    assert not str(payload.get("saved_path", "")).startswith(("C:", "D:", "/"))
    job_response = client.get(
        f"/jobs/{payload['job_id']}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert job_response.status_code == 200
    assert job_response.json()["warranty_id"] == payload["warranty_id"]
    with SessionLocal() as db:
        job = db.query(PipelineJobDB).filter_by(id=payload["job_id"]).first()
        assert job is not None


def test_pipeline_completes_with_pdf(tmp_path):
    pdf_path = tmp_path / "invoice.pdf"
    _make_pdf(pdf_path, "Brand: Acmeco Model: ZX-100 Purchase date: 2025-01-01 Warranty 12 months")
    artifact = ingest_artifact(ArtifactType.invoice, file_path=str(pdf_path), use_ocr=False)
    warranty = canonicalize_artifact(artifact, None)

    with SessionLocal() as db:
        job = invoice_pipeline.create_job(
            db,
            warranty_id=warranty.id,
            artifact_id=artifact.id,
            source_path=str(pdf_path),
        )
    invoice_pipeline.run_job(job.id)
    with SessionLocal() as db:
        job_row = db.query(PipelineJobDB).filter_by(id=job.id).first()
        assert job_row is not None
        assert job_row.status == "done"
        summary = db.query(WarrantySummaryDB).filter_by(warranty_id=warranty.id).first()
        assert summary is not None


def test_pipeline_with_mock_text():
    text = "Invoice No: INV-123 Brand: Acmeco Model: ZX-100 Purchase date: 2025-01-01"
    artifact = ingest_artifact(ArtifactType.invoice, content=text, use_ocr=False)
    warranty = canonicalize_artifact(artifact, None)
    with SessionLocal() as db:
        job = invoice_pipeline.create_job(
            db,
            warranty_id=warranty.id,
            artifact_id=artifact.id,
            source_path=None,
        )
    invoice_pipeline.run_job(job.id)
    with SessionLocal() as db:
        parsed = (
            db.query(ParsedFieldDB)
            .filter_by(warranty_id=warranty.id)
            .order_by(ParsedFieldDB.created_at.desc())
            .first()
        )
        assert parsed is not None
        assert parsed.brand == "Acmeco"


def test_pipeline_completes_with_docx(tmp_path):
    from docx import Document

    docx_path = tmp_path / "invoice.docx"
    doc = Document()
    doc.add_paragraph("Invoice No: INV-456")
    doc.add_paragraph("Brand: Acmeco")
    doc.add_paragraph("Model: ZX-200")
    doc.add_paragraph("Purchase date: 2025-01-01")
    doc.add_paragraph("Warranty 24 months")
    doc.save(str(docx_path))

    artifact = ingest_artifact(ArtifactType.invoice, file_path=str(docx_path), use_ocr=True)
    assert "Acmeco" in (artifact.content or "")
    assert "[OCR note]" not in (artifact.content or "")

    warranty = canonicalize_artifact(artifact, None)
    with SessionLocal() as db:
        job = invoice_pipeline.create_job(
            db,
            warranty_id=warranty.id,
            artifact_id=artifact.id,
            source_path=str(docx_path),
        )
    invoice_pipeline.run_job(job.id)
    with SessionLocal() as db:
        job_row = db.query(PipelineJobDB).filter_by(id=job.id).first()
        assert job_row is not None
        assert job_row.status == "done"
        summary = db.query(WarrantySummaryDB).filter_by(warranty_id=warranty.id).first()
        assert summary is not None


def test_summary_template_when_llm_disabled():
    summary_engine._LLM_PROVIDER = "none"
    warranty = CanonicalWarranty(
        id="wty_test",
        brand="Acmeco",
        model_code="ZX-100",
        coverage_months=12,
        terms=["Coverage applies under normal usage."],
        exclusions=["Physical damage excluded."],
        claim_steps=["Keep invoice ready."],
    )
    text, source = summary_engine.summarize_warranty(warranty)
    assert source == "template"
    assert "Coverage" in text


def test_openai_invoice_enrichment_does_not_override_high_confidence_fields():
    fields = {"brand": "Acmeco", "model_code": "ZX-100"}
    confidence = {"brand": 0.9, "model_code": 0.8}
    enrichment = {
        "fields": {"brand": "WrongCo", "product_category": "Microwave"},
        "confidence": {"brand": 0.85, "product_category": 0.7},
        "model": "gpt-test",
        "reasoning": "visible in invoice",
    }
    merged_fields, merged_confidence, meta = merge_invoice_enrichment(fields, confidence, enrichment)
    assert merged_fields["brand"] == "Acmeco"
    assert merged_fields["product_category"] == "Microwave"
    assert merged_confidence["brand"] == 0.9
    assert "product_category" in meta["fields"]
    assert "brand" not in meta["fields"]


def test_openai_summary_provider_falls_back_without_text(monkeypatch):
    monkeypatch.setattr(summary_engine, "_LLM_PROVIDER", "openai")
    monkeypatch.setattr(summary_engine, "_OPENAI_FALLBACK_PROVIDER", "template")
    monkeypatch.setattr(summary_engine, "_summarize_with_openai", lambda prompt: (None, "disabled"))
    warranty = CanonicalWarranty(
        id="wty_test_openai",
        brand="Acmeco",
        model_code="ZX-100",
        coverage_months=12,
        terms=["Coverage applies under normal usage."],
    )
    text, source = summary_engine.summarize_warranty(warranty)
    assert source == "template"
    assert "Coverage" in text


def test_invoice_parser_prefers_line_item_oem_over_seller_header():
    text = """
TAX INVOICE
The Print Mall Invoice No. Dated
SHOP NO-1,CAPRI TRADE CENTRE TPM/4313/25-26 1-Jul-25
Si Description of Goods HSNISAC Quantity Rate Amount
1 Epson L 3250 Printer 84433240 1no 13,200.00 11,186.44
XAHT699208
2 Paper Rim 48025890 1no 250.00 211.86
Warranty services and claims, if any to be settled and borne by the manufactures.
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert fields["brand"] == "Epson"
    assert fields["model_code"] == "L3250"
    assert fields["serial_no"] == "XAHT699208"
    assert fields["product_category"] == "electronics"
    assert alternatives["seller"] == ["The Print Mall"]
    assert "Epson L 3250 Printer" in alternatives["product_line"][0]
    assert confidence["brand"] >= 0.8


def test_invoice_parser_rejects_recipient_header_and_spec_as_identity():
    text = """
Tax Invoice
Original For Recipient
Bill To: Gaurav Customer
Description | Refresh Rate | Monster 6000 mAh Battery | IP54 | 6
1 Monster 6000 mAh Battery IP54 1 pcs 1,499.00
Invoice Date: 01-05-2026
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert "brand" not in fields
    assert fields.get("brand") not in {"Original For Recipient", "(Original For Recipient)", "Recipient"}
    assert fields.get("model_code") not in {"IP54", "6", "ORIGINAL"}
    assert fields.get("product_name")
    assert fields["product_name"] == "Monster 6000 mAh Battery"
    assert fields["purchase_date"] == "2026-05-01"
    assert "discarded_identity_candidates" not in alternatives or all(
        "Monster 6000 mAh Battery" not in item for item in alternatives["discarded_identity_candidates"]
    )


def test_invoice_parser_handles_samsung_mobile_pipe_specs_without_fake_serial():
    text = """
Tax Invoice
Order Date: 01.05.2026
Invoice Number : DEL5-53804
Invoice Date : 02.05.2026
Sl. No Description Unit Price Discount Qty Net Amount Tax Rate Tax Type Tax Amount Total Amount
1 Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage) | Smoothest 120 Hz Refresh Rate| Monster 6000 mAh Battery | IP54 | 6 Gen OS Upgrades | AI | Gemini Live | Without Charger | B0GN1NNYXF ( SMNG-M17e-VIOLET-6+128GB ) HSN:85171300
Shipping Charges
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert fields["brand"] == "Samsung"
    assert fields["model_code"] == "M17E"
    assert fields["product_name"] == "Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage)"
    assert fields["purchase_date"] == "2026-05-01"
    assert fields["invoice_no"] == "DEL5-53804"
    assert fields.get("serial_no") != "85171300"
    assert "IP54" not in fields.values()
    assert alternatives["product_line"]


def test_invoice_parser_handles_wrapped_samsung_mobile_ocr_lines():
    text = """
Tax Invoice
Order Number: 171-8910655-3956348
Order Date: 01.05.2026
Invoice Number : DEL5-53804
Invoice Details : HR-DEL5-1224631255-2627
Invoice Date : 02.05.2026
Sl. No Description Unit Price Discount Qty Net Amount Tax Rate Tax Type Tax Amount Total Amount
1 Samsung Galaxy M17e 5G Mobile (Vibe Violet,
6GB RAM, 128GB Storage) | Smoothest 120 Hz
Refresh Rate| Monster 6000 mAh Battery | IP54 | 6
Gen OS Upgrades | AI | Gemini Live | Without
Charger | B0GN1NNYXF (
SMNG-M17e-VIOLET-6+128GB )
HSN:85171300
Shipping Charges
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert fields["brand"] == "Samsung"
    assert fields["model_code"] == "M17E"
    assert fields["product_name"] == "Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage)"
    assert fields["purchase_date"] == "2026-05-01"
    assert fields["invoice_no"] == "DEL5-53804"
    assert fields["product_category"] == "mobile"
    assert fields.get("serial_no") != "85171300"
    assert alternatives.get("seller") != ["Order Number: 171-8910655-3956348"]
    assert "Samsung Galaxy M17e 5G Mobile" in alternatives["product_line"][0]


def test_invoice_parser_handles_standalone_row_marker_wrapped_samsung_mobile():
    text = """
Tax Invoice
Order Number: 171-8910655-3956348
Order Date: 01.05.2026
Invoice Number : DEL5-53804
Invoice Details : HR-DEL5-1224631255-2627
Invoice Date : 02.05.2026
Sl. No Description Unit Price Discount Qty Net Amount Tax Rate Tax Type Tax Amount Total Amount
1
Samsung Galaxy M17e 5G Mobile (Vibe Violet,
6GB RAM, 128GB Storage) | Smoothest 120 Hz
Refresh Rate| Monster 6000 mAh Battery | IP54 | 6
Gen OS Upgrades | AI | Gemini Live | Without
Charger | B0GN1NNYXF (
SMNG-M17e-VIOLET-6+128GB )
HSN:85171300
Shipping Charges
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert fields["brand"] == "Samsung"
    assert fields["model_code"] == "M17E"
    assert fields["product_name"] == "Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage)"
    assert fields["purchase_date"] == "2026-05-01"
    assert fields["invoice_no"] == "DEL5-53804"
    assert fields["product_category"] == "mobile"
    assert fields.get("serial_no") != "85171300"
    assert "Samsung Galaxy M17e 5G Mobile" in alternatives["product_line"][0]


def test_invoice_parser_rejects_address_block_and_extracts_later_wrapped_samsung_item():
    text = """
Tax Invoice
Sold By : Sattva Horizon Private Limited
Billing Address :
Gaurav
Reet/Killa Nos. 38//8/2 min, 192//22/1, 196//2/1/1,
37//15/1, 15/2,, Adjacent to Starex School, Village
- Binola, National Highway -8, Tehsil - Manesar
Gurgaon, Haryana, 122413
Shipping Address :
Gaurav
79/1nashvilla road, Near dastaana factory
DEHRADUN, UTTARAKHAND, 248001
Place of supply: UTTARAKHAND
Place of delivery: UTTARAKHAND
Order Number: 171-8910655-3956348
Order Date: 01.05.2026
Invoice Number : DEL5-53804
Invoice Details : HR-DEL5-1224631255-2627
Invoice Date : 02.05.2026
Sl. No Description Unit Price Discount Qty Net Amount Tax Rate Tax Type Tax Amount Total Amount
1 Samsung Galaxy M17e 5G Mobile (Vibe Violet,
6GB RAM, 128GB Storage) | Smoothest 120 Hz
Refresh Rate| Monster 6000 mAh Battery | IP54 | 6
Gen OS Upgrades | AI | Gemini Live | Without
Charger | B0GN1NNYXF (
SMNG-M17e-VIOLET-6+128GB )
HSN:85171300
Shipping Charges
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert fields["brand"] == "Samsung"
    assert fields["model_code"] == "M17E"
    assert fields["product_name"] == "Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage)"
    assert fields["purchase_date"] == "2026-05-01"
    assert fields["invoice_no"] == "DEL5-53804"
    assert fields["product_category"] == "mobile"
    assert fields.get("serial_no") != "85171300"
    assert "Sattva" not in fields["product_name"]
    assert "Survey" not in fields["product_name"]
    assert "Samsung Galaxy M17e 5G Mobile" in alternatives["product_line"][0]


def test_invoice_parser_strips_merged_table_header_before_samsung_item():
    text = """
Tax Invoice
Invoice Date : 02.05.2026
Sl. No Description Unit Price Discount Qty Net Amount Tax Rate Tax Type Tax Amount Total Amount 1 Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage) | Smoothest 120 Hz Refresh Rate| Monster 6000 mAh Battery | IP54 | 6 Gen OS Upgrades | AI | Gemini Live | Without Charger | B0GN1NNYXF ( SMNG-M17e-VIOLET-6+128GB ) HSN:85171300
"""
    fields, confidence, alternatives = extract_product_fields(text)

    assert fields["brand"] == "Samsung"
    assert fields["model_code"] == "M17E"
    assert fields["product_name"] == "Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage)"
    assert "Description Unit Price" not in alternatives["product_line"][0]


def test_invoice_parser_ignores_ocr_note_file_paths_as_product_identity():
    text = "[OCR note] File not found: C:\\Users\\lenovo\\Desktop\\invoice.pdf"
    fields, confidence, alternatives = extract_product_fields(text)

    assert "brand" not in fields
    assert "product_name" not in fields
    assert "model_code" not in fields
    assert alternatives["notes"] == ["No strong signals found; manual entry may be required."]


def test_invoice_enrichment_sanitizer_removes_bad_ai_identity_fields():
    fields = {"product_name": "Monster 6000 mAh Battery"}
    confidence = {"product_name": 0.75}
    enrichment = {
        "fields": {"brand": "(Original For Recipient)", "model_code": "IP54"},
        "confidence": {"brand": 0.85, "model_code": 0.85},
    }

    merged_fields, merged_confidence, _ = merge_invoice_enrichment(fields, confidence, enrichment)
    sanitized_fields, sanitized_confidence, alternatives = invoice_pipeline.sanitize_invoice_identity_fields(
        merged_fields,
        merged_confidence,
        {},
    )

    assert sanitized_fields == {"product_name": "Monster 6000 mAh Battery"}
    assert "brand" not in sanitized_confidence
    assert "model_code" not in sanitized_confidence
    assert alternatives["discarded_identity_candidates"]


def test_canonical_invoice_does_not_invent_unconfirmed_warranty_terms():
    artifact = ingest_artifact(
        ArtifactType.invoice,
        content="Tax Invoice\nSeller Store\n1 Epson L 3250 Printer 84433240 1no\nXAHT699208\nDated 1-Jul-25",
    )
    warranty = canonicalize_artifact(artifact, None)

    assert warranty.brand == "Epson"
    assert warranty.terms == []
    assert warranty.exclusions == []
    assert any("Verify official OEM warranty terms" in step for step in warranty.claim_steps)


def test_epson_l3250_discovers_official_source_and_terms():
    sources = discover_sources(
        brand="Epson",
        model_code="L3250",
        product_name="Epson L 3250 Printer",
        region="IN",
        mode="auto+manual",
        allow_retail=True,
    )

    assert sources
    assert sources[0].official is True
    assert "epson.co.in" in sources[0].url

    with SessionLocal() as db:
        result = lookup_terms(
            db,
            brand="Epson",
            category="electronics",
            region="IN",
            model_code="L3250",
            product_name="Epson L 3250 Printer",
            force_refresh=True,
        )

    assert result.duration_months == 12
    assert result.source_url and "epson.co.in" in result.source_url


def test_terms_lookup_merges_multiple_controlled_oem_sources(monkeypatch):
    urls = [
        "https://support.acmeco.example/product/zx-100",
        "https://support.acmeco.example/warranty",
        "https://support.acmeco.example/claim",
    ]

    def fake_discover_sources(**kwargs):
        return [
            DiscoverySource(url=urls[0], source_type="oem_product", score=90, official=True),
            DiscoverySource(url=urls[1], source_type="oem_warranty", score=85, official=True),
            DiscoverySource(url=urls[2], source_type="oem_warranty", score=80, official=True),
        ]

    def fake_parse_terms_from_url(url):
        if url.endswith("/product/zx-100"):
            return ParsedTerms(
                duration_months=12,
                terms=["Standard coverage for 12 months from purchase date."],
                exclusions=[],
                claim_steps=[],
                raw_text="ZX-100 product page",
            ), None
        if url.endswith("/warranty"):
            return ParsedTerms(
                duration_months=None,
                terms=[],
                exclusions=["Liquid damage is excluded."],
                claim_steps=[],
                raw_text="Warranty policy page",
            ), None
        return ParsedTerms(
            duration_months=None,
            terms=[],
            exclusions=[],
            claim_steps=["Keep invoice and serial number ready."],
            raw_text="Claim support page",
        ), None

    monkeypatch.setattr(terms_lookup, "discover_sources", fake_discover_sources)
    monkeypatch.setattr(terms_lookup, "parse_terms_from_url", fake_parse_terms_from_url)

    with SessionLocal() as db:
        result = lookup_terms(
            db,
            brand="Acmeco",
            category="electronics",
            region="US",
            model_code="ZX-100",
            product_name="Acmeco ZX-100 Printer",
            force_refresh=True,
        )

    assert result.duration_months == 12
    assert "Standard coverage for 12 months from purchase date." in result.terms
    assert "Liquid damage is excluded." in result.exclusions
    assert "Keep invoice and serial number ready." in result.claim_steps
    assert result.source_urls == urls


def test_terms_lookup_normalizes_samsung_mobile_official_page(monkeypatch):
    samsung_url = "https://www.samsung.com/in/support/warranty/"

    def fake_discover_sources(**kwargs):
        return [
            DiscoverySource(
                url=samsung_url,
                source_type="oem_warranty",
                score=95,
                official=True,
                brand="Samsung",
                region="IN",
            )
        ]

    def fake_parse_terms_from_url(url):
        assert url == samsung_url
        return ParsedTerms(
            duration_months=60,
            terms=[
                "Standard coverage for 60 months from purchase date.",
                "Limited International One Year Warranty",
                "Warranty does not cover normal wear and tear, including batteries or displays.",
            ],
            exclusions=[
                "Warranty does not cover repair due to external factors.",
            ],
            claim_steps=[
                "News & Alerts",
                "Warranty Check",
                "Check Repair Status",
            ],
            raw_text=(
                "Samsung support warranty. Limited International One Year Warranty. "
                "The limited warranty period of 1 year applies to mobile products."
            ),
        ), None

    monkeypatch.setattr(terms_lookup, "discover_sources", fake_discover_sources)
    monkeypatch.setattr(terms_lookup, "parse_terms_from_url", fake_parse_terms_from_url)

    with SessionLocal() as db:
        result = lookup_terms(
            db,
            brand="Samsung",
            category="mobile",
            region="IN",
            model_code="M17E",
            product_name="Samsung Galaxy M17e 5G Mobile",
            force_refresh=True,
        )

    assert result.duration_months == 12
    assert not any("60 months" in term for term in result.terms)
    assert "Limited International One Year Warranty" in result.terms
    assert "Warranty Check" in result.claim_steps
    assert "News & Alerts" not in result.claim_steps


def test_warranty_list_includes_invoice_and_upload_display_label():
    warranty_id = "wty_label_test"
    with SessionLocal() as db:
        db.query(ParsedFieldDB).filter_by(warranty_id=warranty_id).delete()
        db.query(WarrantyDB).filter_by(id=warranty_id).delete()
        warranty = WarrantyDB(
            id=warranty_id,
            product_name="Samsung Galaxy M17e 5G Mobile",
            brand="Samsung",
            model_code="M17E",
            purchase_date=datetime(2026, 5, 1),
            coverage_months=12,
            expiry_date=datetime(2027, 5, 1),
            created_at=datetime(2026, 5, 2, 10, 30),
        )
        parsed = ParsedFieldDB(
            warranty_id=warranty_id,
            brand="Samsung",
            model_code="M17E",
            product_name="Samsung Galaxy M17e 5G Mobile",
            invoice_no="DEL5-53804",
            purchase_date=datetime(2026, 5, 1),
            created_at=datetime(2026, 5, 2, 10, 31),
        )
        db.add(warranty)
        db.add(parsed)
        db.commit()

    client = TestClient(app)
    login = client.post(
        "/auth/login",
        data={"username": "admin", "password": "admin123"},
        headers={"accept": "application/json"},
    )
    assert login.status_code == 200
    res = client.get("/warranties/list")
    assert res.status_code == 200
    item = next(w for w in res.json()["warranties"] if w["id"] == warranty_id)
    assert item["invoice_no"] == "DEL5-53804"
    assert "Samsung Galaxy M17e 5G Mobile" in item["display_label"]
    assert "Inv DEL5-53804" in item["display_label"]
    assert "Bought 2026-05-01" in item["display_label"]
    assert "Uploaded 2026-05-02 10:30" in item["display_label"]


def test_warranty_list_prefers_parsed_product_when_warranty_is_placeholder():
    warranty_id = "wty_label_placeholder"
    with SessionLocal() as db:
        db.query(ParsedFieldDB).filter_by(warranty_id=warranty_id).delete()
        db.query(WarrantyDB).filter_by(id=warranty_id).delete()
        warranty = WarrantyDB(
            id=warranty_id,
            product_name="Product",
            brand=None,
            model_code=None,
            purchase_date=datetime(2026, 5, 1),
            coverage_months=12,
            expiry_date=datetime(2027, 5, 1),
            created_at=datetime(2026, 5, 2, 10, 30),
        )
        parsed = ParsedFieldDB(
            warranty_id=warranty_id,
            brand="Samsung",
            model_code="M17E",
            product_name="Samsung Galaxy M17e 5G Mobile",
            invoice_no="DEL5-53804",
            purchase_date=datetime(2026, 5, 1),
            created_at=datetime(2026, 5, 2, 10, 31),
        )
        db.add(warranty)
        db.add(parsed)
        db.commit()

    client = TestClient(app)
    login = client.post(
        "/auth/login",
        data={"username": "admin", "password": "admin123"},
        headers={"accept": "application/json"},
    )
    assert login.status_code == 200
    res = client.get("/warranties/list")
    assert res.status_code == 200
    item = next(w for w in res.json()["warranties"] if w["id"] == warranty_id)
    assert item["display_label"].startswith("Samsung Galaxy M17e 5G Mobile | Inv DEL5-53804")
    assert "Uploaded 2026-05-02 10:30" in item["display_label"]


def test_pipeline_persists_epson_terms_source_after_lookup():
    artifact = ingest_artifact(
        ArtifactType.invoice,
        content=(
            "TAX INVOICE\n"
            "The Print Mall Invoice No. Dated\n"
            "TPM/4313/25-26 1-Jul-25\n"
            "1 Epson L 3250 Printer 84433240 1no 13,200.00 11,186.44\n"
            "XAHT699208\n"
        ),
    )
    warranty = canonicalize_artifact(artifact, None)
    with SessionLocal() as db:
        job = invoice_pipeline.create_job(
            db,
            warranty_id=warranty.id,
            artifact_id=artifact.id,
            source_path=None,
        )

    invoice_pipeline.run_job(job.id)

    with SessionLocal() as db:
        job_row = db.query(PipelineJobDB).filter_by(id=job.id).first()
        warranty_row = db.query(WarrantyDB).filter_by(id=warranty.id).first()
        summary = db.query(WarrantySummaryDB).filter_by(warranty_id=warranty.id).first()

    assert job_row.status == "done"
    assert job_row.error is None
    assert warranty_row.brand == "Epson"
    assert warranty_row.model_code == "L3250"
    assert warranty_row.coverage_months == 12
    assert warranty_row.expiry_date is not None
    assert "epson.co.in" in (warranty_row.alternatives or {}).get("terms_source_url", "")
    assert any(
        "epson.co.in" in url
        for url in (warranty_row.alternatives or {}).get("terms_source_urls", [])
    )
    assert (warranty_row.alternatives or {}).get("terms_source_type") == "approved_oem_source"
    assert summary is not None


def test_pipeline_force_refreshes_terms_for_new_upload(monkeypatch):
    calls = []

    def fake_lookup_terms(*args, **kwargs):
        calls.append(kwargs)
        return TermsResult(
            duration_months=12,
            terms=["Fresh official warranty terms from parser."],
            source_url="https://www.epson.co.in/fresh",
            source_urls=["https://www.epson.co.in/fresh"],
        )

    monkeypatch.setattr(invoice_pipeline, "lookup_terms", fake_lookup_terms)

    artifact = ingest_artifact(
        ArtifactType.invoice,
        content=(
            "TAX INVOICE\n"
            "The Print Mall Invoice No. Dated\n"
            "TPM/4313/25-26 1-Jul-25\n"
            "1 Epson L 3250 Printer 84433240 1no 13,200.00 11,186.44\n"
            "XAHT699208\n"
        ),
    )
    warranty = canonicalize_artifact(artifact, None)
    with SessionLocal() as db:
        job = invoice_pipeline.create_job(
            db,
            warranty_id=warranty.id,
            artifact_id=artifact.id,
            source_path=None,
        )

    invoice_pipeline.run_job(job.id)

    assert calls
    assert calls[0]["force_refresh"] is True


def test_terms_lookup_runs_when_identity_exists_even_with_default_coverage():
    warranty = WarrantyDB(
        id="wty_lookup_identity",
        product_name="Samsung Galaxy M17e 5G Mobile",
        brand="Samsung",
        model_code="M17E",
        coverage_months=12,
    )
    placeholder = WarrantyDB(
        id="wty_lookup_placeholder",
        product_name="Product",
        brand=None,
        model_code=None,
        coverage_months=12,
    )

    assert invoice_pipeline._has_terms_lookup_identity(warranty, {}) is True
    assert invoice_pipeline._has_terms_lookup_identity(placeholder, {}) is False


def test_pipeline_summary_upstream_error_does_not_fail_upload(monkeypatch):
    def fake_lookup_terms(*args, **kwargs):
        return None

    def fail_summary(*args, **kwargs):
        raise RuntimeError("upstream error")

    monkeypatch.setattr(invoice_pipeline, "lookup_terms", fake_lookup_terms)
    monkeypatch.setattr(invoice_pipeline, "summarize_warranty", fail_summary)
    monkeypatch.setattr(invoice_pipeline, "build_structured_summary", fail_summary)

    artifact = ingest_artifact(
        ArtifactType.invoice,
        content=(
            "Tax Invoice\n"
            "Order Date: 01.05.2026\n"
            "Invoice Number : DEL5-53804\n"
            "Invoice Date : 02.05.2026\n"
            "1 Samsung Galaxy M17e 5G Mobile (Vibe Violet, 6GB RAM, 128GB Storage)\n"
        ),
    )
    warranty = canonicalize_artifact(artifact, None)
    with SessionLocal() as db:
        job = invoice_pipeline.create_job(
            db,
            warranty_id=warranty.id,
            artifact_id=artifact.id,
            source_path=None,
        )

    invoice_pipeline.run_job(job.id)

    with SessionLocal() as db:
        job_row = db.query(PipelineJobDB).filter_by(id=job.id).first()
        summary = db.query(WarrantySummaryDB).filter_by(warranty_id=warranty.id).first()

    assert job_row.status == "done"
    assert job_row.error is None
    assert summary is not None
    assert summary.source == "template_fallback"
